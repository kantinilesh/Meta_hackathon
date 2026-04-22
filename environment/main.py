from fastapi import Body, FastAPI, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
import asyncio, uuid, os, secrets, logging, time, json, re
from typing import Dict, List, Optional
from openai import AsyncOpenAI
import httpx
from pathlib import Path

from .models import (
    NegotiationSession, PartyConfig, SessionStatus, Observation, 
    Action, Reward, TaskConfig, GradeResult, NegotiationTurn, NegotiationRole,
    CompanyDocument
)
from .env import ContractEnv
from .dual_env import DualAgentEnv
from .agent_runner import AgentRunner
from .contracts.nda_template import TASK_CONTRACTS, load_contract
from .contracts.product_sales_template import load_product_contract

app = FastAPI(title="ContractEnv", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

UPLOAD_DIR = Path(os.path.join(os.path.dirname(__file__), "..", "uploads"))
UPLOAD_DIR.mkdir(exist_ok=True)

env_sessions: Dict[str, ContractEnv] = {}
negotiation_sessions: Dict[str, NegotiationSession] = {}
ws_connections: Dict[str, List[WebSocket]] = {}
session_documents: Dict[str, List[CompanyDocument]] = {}

import os
from dotenv import load_dotenv
load_dotenv()
try:
    openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY", os.getenv("API_KEY", "dummy")), base_url=os.getenv("API_BASE_URL", "https://api.openai.com/v1"))
except:
    openai_client = None

# STANDARD OPENENV ENDPOINTS
@app.get("/health")
def health():
    return {"status": "ok", "version": "1.0.0", "timestamp": time.time(), "active_sessions": len(negotiation_sessions)}

class UploadDocumentReq(BaseModel):
    session_id: str
    document_type: str
    extract_terms: bool = True

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".xlsx"}
TEXT_EXTENSIONS = {".txt", ".docx"}

async def _extract_text_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    content = ""
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    content += page.extract_text() or ""
    except Exception as e:
        logging.warning(f"Failed to extract text from {file_path}: {e}")
    return content[:15000]

async def _extract_key_terms(text: str, doc_type: str) -> List[str]:
    if not openai_client:
        return []
    term_prompts = {
        "financials": "Extract 5-10 key financial metrics and figures (revenue, profit, liabilities, assets, valuation) from this document.",
        "bylaws": "Extract 5-10 key governance terms (board size, voting rights, decision thresholds) from this document.",
        "due_diligence": "Extract 5-10 key findings about the target company from this due diligence report.",
        "cap_table": "Extract 5-10 key equity terms (shareholders, ownership percentages, vesting schedules) from this cap table.",
        "employment": "Extract 5-10 key employment terms from this document.",
        "ip_assignment": "Extract 5-10 key IP assignment terms from this document.",
        "other": "Extract 5-10 key terms relevant to M&A negotiation from this document."
    }
    prompt = term_prompts.get(doc_type, term_prompts["other"]) + f"\n\nDocument:\n{text[:8000]}"
    try:
        resp = await openai_client.chat.completions.create(
            model=os.getenv("MODEL_NAME", "gpt-4o-mini"),
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500
        )
        result = resp.choices[0].message.content
        terms = [t.strip() for t in re.findall(r"[\d\.\$%]+|[A-Z][A-Za-z\s]{3,30}", result)][:10]
        return terms
    except Exception as e:
        logging.warning(f"Failed to extract terms: {e}")
        return []

async def _summarize_document(text: str, doc_type: str) -> str:
    if not openai_client:
        return text[:500]
    summary_prompts = {
        "financials": "Summarize the key financial position in 3-5 sentences: total revenue, profitability, key liabilities, and any notable financial events.",
        "bylaws": "Summarize the key governance structure in 3-5 sentences: board composition, voting procedures, and decision-making thresholds.",
        "due_diligence": "Summarize the key findings in 3-5 sentences: company strengths, risks identified, and any deal-breakers.",
        "cap_table": "Summarize the equity structure in 3-5 sentences: major shareholders, ownership percentages, and any vesting conditions.",
        "employment": "Summarize key employment terms in 3-5 sentences: key personnel, contracts, and any golden parachute clauses.",
        "ip_assignment": "Summarize IP ownership in 3-5 sentences: what IP is being assigned, any third-party claims, and exclusivity terms.",
        "other": "Summarize the key contents of this document in 3-5 sentences for M&A negotiation purposes."
    }
    prompt = summary_prompts.get(doc_type, summary_prompts["other"]) + f"\n\nDocument:\n{text[:8000]}"
    try:
        resp = await openai_client.chat.completions.create(
            model=os.getenv("MODEL_NAME", "gpt-4o-mini"),
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300
        )
        return resp.choices[0].message.content
    except Exception as e:
        logging.warning(f"Failed to summarize: {e}")
        return text[:200]

@app.post("/document/upload")
async def upload_document(
    file: UploadFile = File(...),
    session_id: str = "",
    document_type: str = "other",
    extract_terms: bool = True
):
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"File type not allowed. Allowed: {list(ALLOWED_EXTENSIONS)}")
    
    doc_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{doc_id}{file_ext}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    text_content = ""
    if file_ext in TEXT_EXTENSIONS:
        text_content = await _extract_text_from_file(file_path)
    
    summary = await _summarize_document(text_content, document_type) if text_content else ""
    key_terms = await _extract_key_terms(text_content, document_type) if extract_terms and text_content else []
    
    doc = CompanyDocument(
        document_id=doc_id,
        file_name=file.filename,
        file_type=file_ext[1:],
        document_type=document_type,
        summary=summary,
        key_terms=key_terms,
        upload_date=str(time.time())
    )
    
    if session_id:
        if session_id not in session_documents:
            session_documents[session_id] = []
        session_documents[session_id].append(doc)
    
    return {
        "document_id": doc_id,
        "file_name": file.filename,
        "document_type": document_type,
        "summary": summary,
        "key_terms": key_terms,
        "upload_date": doc.upload_date
    }

@app.get("/session/documents")
def get_session_documents(session_id: str):
    return {"documents": session_documents.get(session_id, [])}

class ResetReq(BaseModel):
    task_id: str = "task1"
    session_id: Optional[str] = None

@app.post("/reset")
def reset(req: Optional[ResetReq] = Body(default=None)):
    task_id = req.task_id if req and req.task_id else "task1"
    session_id = req.session_id if req and req.session_id else str(uuid.uuid4())
    env = ContractEnv(task_id=task_id)
    obs = env.reset()
    env_sessions[session_id] = env
    return {"observation": obs.model_dump(), "session_id": session_id}

class StepReq(BaseModel):
    session_id: str
    action: Action

@app.post("/step")
def step(req: StepReq):
    if req.session_id not in env_sessions:
        raise HTTPException(404, "Session not found")
    obs, reward, done, info = env_sessions[req.session_id].step(req.action)
    return {"observation": obs.model_dump(), "reward": reward.model_dump(), "done": done, "info": info}

@app.get("/state")
def state(session_id: str):
    if session_id not in env_sessions:
        raise HTTPException(404, "Session not found")
    return env_sessions[session_id].state()

@app.get("/tasks")
def tasks():
    return [
        TaskConfig(task_id="task1", name="Clause Identification", difficulty="easy", description="", max_turns=20, target_score=0.85),
        TaskConfig(task_id="task2", name="Clause Redlining", difficulty="medium", description="", max_turns=30, target_score=0.65),
        TaskConfig(task_id="task3", name="Full Negotiation", difficulty="hard", description="", max_turns=50, target_score=0.45),
    ]

# SESSION PRODUCT ENDPOINTS
class SessionCreateReq(BaseModel):
    contract_id: str = "nda_001"
    seller_company_name: str
    seller_constraints: List[dict]
    seller_agent_style: str = "balanced"
    seller_context: str = ""

@app.post("/session/create")
def session_create(req: SessionCreateReq):
    session_id = str(uuid.uuid4())
    invite_token = secrets.token_urlsafe(16)
    
    if req.contract_id == "product_001":
        contract_data = load_product_contract()
        req.seller_constraints.append({
            "constraint_id": "sys_spec",
            "description": "The Product Specifications MUST NOT be modified from the original text.",
            "clause_category": "scope",
            "is_deal_breaker": True,
            "rule_type": "must_include",
            "rule_value": "original text",
            "priority": 10
        })
    else:
        contract_data = load_contract("task3") # Base contract
    
    seller_config = PartyConfig(
        role="seller",
        company_name=req.seller_company_name,
        constraints=req.seller_constraints,
        agent_style=req.seller_agent_style,
        constraint_summary="Custom constraints",
        company_context=req.seller_context
    )
    
    session = NegotiationSession(
        session_id=session_id,
        status="waiting_client",
        seller_config=seller_config,
        contract_id=req.contract_id,
        contract_title=contract_data["title"],
        clauses=contract_data["clauses"],
        created_at=str(time.time()),
        invite_token=invite_token
    )
    negotiation_sessions[session_id] = session
    
    return {
        "session_id": session_id,
        "invite_token": invite_token,
        "invite_url": f"http://localhost:3000/join/{invite_token}",
        "status": session.status
    }

class SessionJoinReq(BaseModel):
    invite_token: str
    client_company_name: str
    client_constraints: List[dict]
    client_agent_style: str = "balanced"
    client_context: str = ""

@app.post("/session/join")
def session_join(req: SessionJoinReq):
    session = next((s for s in negotiation_sessions.values() if s.invite_token == req.invite_token), None)
    if not session:
        raise HTTPException(404, "Invalid token")
        
    session.client_config = PartyConfig(
        role="client",
        company_name=req.client_company_name,
        constraints=req.client_constraints,
        agent_style=req.client_agent_style,
        constraint_summary="Custom constraints",
        company_context=req.client_context
    )
    session.status = "ready"
    return {"session_id": session.session_id, "status": session.status}

class GradeReq(BaseModel):
    session_id: str
    task_id: str


def _strict_unit_interval(value: float) -> float:
    return max(0.01, min(0.99, value))


@app.post("/grade")
def grade(req: GradeReq):
    env = env_sessions.get(req.session_id)
    if not env:
        raise HTTPException(404, "Session not found")

    raw_score = sum(t.reward_delta for t in env.history if getattr(t, "reward_delta", None)) if env.history else 0.0

    if req.task_id == "task3":
        total_clauses = len(env.clauses) or 1
        agreed_clauses = sum(1 for clause in env.clauses.values() if clause.status == "agreed")
        score = _strict_unit_interval(max(raw_score, agreed_clauses / total_clauses))
        passed = agreed_clauses == total_clauses
        details = [
            f"Negotiation reached agreement on {agreed_clauses}/{total_clauses} clauses."
        ]
    elif req.task_id == "task2":
        max_task_score = max(0.28 * len(env.clauses), 0.28)
        score = _strict_unit_interval(raw_score / max_task_score)
        passed = raw_score > 0
        details = ["Redlining quality normalized against the task's clause budget."]
    elif req.task_id == "task1":
        unfair_clause_count = sum(1 for clause in env.clauses.values() if clause.ground_truth_label == "unfair") or 1
        max_task_score = 0.15 * unfair_clause_count
        score = _strict_unit_interval(raw_score / max_task_score)
        passed = raw_score > 0
        details = ["Clause-labeling score normalized against the number of unfair clauses."]
    else:
        score = _strict_unit_interval(raw_score)
        passed = raw_score > 0
        details = ["Grading complete based on accumulated reward."]
    
    return {
        "task_id": req.task_id, 
        "score": score, 
        "breakdown": {}, 
        "passed": passed, 
        "details": details
    }

class LinkDocumentReq(BaseModel):
    session_id: str
    role: str
    document_ids: List[str]

@app.post("/session/link-documents")
def link_documents(req: LinkDocumentReq):
    if req.session_id not in session_documents:
        raise HTTPException(404, "No documents found for session")
    
    docs = session_documents[req.session_id]
    if req.session_id in negotiation_sessions:
        session = negotiation_sessions[req.session_id]
        if req.role == "seller" and session.seller_config:
            session.seller_config.documents = docs
        elif req.role == "client" and session.client_config:
            session.client_config.documents = docs
    
    return {"linked": len(docs), "documents": [d.model_dump() for d in docs]}

class SessionStartReq(BaseModel):
    session_id: str

@app.post("/session/start")
async def session_start(req: SessionStartReq):
    if req.session_id not in negotiation_sessions:
        raise HTTPException(404, "Session not found")
    session = negotiation_sessions[req.session_id]
    session.status = "negotiating"
    
    # Launch background negotiation
    asyncio.create_task(run_negotiation(req.session_id))
    
    return {"session_id": req.session_id, "status": "negotiating", "websocket_url": f"ws://localhost:7860/ws/{req.session_id}"}

class SessionSignReq(BaseModel):
    session_id: str
    role: str

@app.post("/session/sign")
def session_sign(req: SessionSignReq):
    if req.session_id not in negotiation_sessions:
        raise HTTPException(404, "Session not found")
    session = negotiation_sessions[req.session_id]
    if req.role == "seller":
        session.seller_signed = True
    else:
        session.client_signed = True
    return {"status": "signed"}

# ── HUMAN INTERVENTION ──────────────────────────────────────────────

class SessionInterveneReq(BaseModel):
    session_id: str
    role: str  # seller or client
    rewind_to_turn: int = 0  # 0 = restart from scratch
    updated_constraints: List[dict] = []

@app.post("/session/intervene")
async def session_intervene(req: SessionInterveneReq):
    """Pause a live negotiation, rewind history, and update constraints."""
    if req.session_id not in negotiation_sessions:
        raise HTTPException(404, "Session not found")
    session = negotiation_sessions[req.session_id]

    # 1. Pause the negotiation loop
    session.status = "paused"
    await asyncio.sleep(0.5)  # give the loop time to stop

    # 2. Broadcast pause event
    await broadcast_turn(req.session_id, NegotiationTurn(
        turn_number=session.turn, speaker="system", action_type="skip", clause_id="",
        content=f"⏸️ Human intervention by {req.role}. Negotiation paused at turn {session.turn}."
    ))

    # 3. Rewind history to the requested turn
    session.negotiation_history = [
        t for t in session.negotiation_history if t.turn_number <= req.rewind_to_turn
    ]
    session.turn = req.rewind_to_turn

    # 4. Reset clause statuses for clauses that were agreed AFTER the rewind point
    agreed_clause_ids_to_reset = []
    for t in list(session.final_agreed_clauses.keys()):
        # Check if the agreement happened after the rewind point
        agreement_turn = next(
            (h for h in session.negotiation_history
             if h.clause_id == t and h.action_type == "accept"),
            None
        )
        if agreement_turn is None:
            agreed_clause_ids_to_reset.append(t)
    for cid in agreed_clause_ids_to_reset:
        del session.final_agreed_clauses[cid]
        for c in session.clauses:
            if c.id == cid:
                c.status = "pending"
                c.current_proposed_text = None

    # 5. Update constraints if provided
    if req.updated_constraints:
        if req.role == "seller" and session.seller_config:
            session.seller_config.constraints = req.updated_constraints
            session.seller_config.constraint_summary = "Updated constraints (human intervention)"
        elif req.role == "client" and session.client_config:
            session.client_config.constraints = req.updated_constraints
            session.client_config.constraint_summary = "Updated constraints (human intervention)"

    await broadcast_turn(req.session_id, NegotiationTurn(
        turn_number=session.turn, speaker="system", action_type="skip", clause_id="",
        content=f"🔄 Constraints updated by {req.role}. Negotiation rewound to turn {req.rewind_to_turn}. Ready to resume."
    ))

    return {
        "status": "paused",
        "rewound_to": req.rewind_to_turn,
        "constraints_updated": len(req.updated_constraints) > 0,
        "history_length": len(session.negotiation_history)
    }

class SessionResumeReq(BaseModel):
    session_id: str

@app.post("/session/resume")
async def session_resume(req: SessionResumeReq):
    """Resume a paused negotiation after human intervention."""
    if req.session_id not in negotiation_sessions:
        raise HTTPException(404, "Session not found")
    session = negotiation_sessions[req.session_id]
    if session.status != "paused":
        raise HTTPException(400, "Session is not paused")
    session.status = "negotiating"

    await broadcast_turn(req.session_id, NegotiationTurn(
        turn_number=session.turn, speaker="system", action_type="skip", clause_id="",
        content="▶️ Negotiation resumed with updated constraints."
    ))

    asyncio.create_task(run_negotiation(req.session_id))
    return {"status": "negotiating", "resumed_from_turn": session.turn}


@app.get("/session/contract")
def session_contract(session_id: str):
    if session_id not in negotiation_sessions:
        raise HTTPException(404, "Session not found")
    session = negotiation_sessions[session_id]
    
    # Rebuild final contract text
    final_text = ""
    for c in session.clauses:
        final_text += f"{c.title}:\n{session.final_agreed_clauses.get(c.id, c.text)}\n\n"
        
    return {"final_contract_text": final_text}
@app.get("/session/status")
def session_status(session_id: str = None, invite_token: str = None, role: str = None):
    session = None
    if session_id:
        session = negotiation_sessions.get(session_id)
    elif invite_token:
        session = next((s for s in negotiation_sessions.values() if s.invite_token == invite_token), None)
        
    if not session:
        raise HTTPException(404, "Session not found")
        
    # Strip opposing party config
    out = session.model_dump()
    if role == "seller":
        out["client_config"] = None
    elif role == "client":
        out["seller_config"] = None
        
    return out

@app.websocket("/ws/{session_id}")
async def ws_endpoint(websocket: WebSocket, session_id: str, role: str = "both"):
    await websocket.accept()
    if session_id not in ws_connections:
        ws_connections[session_id] = []
    ws_connections[session_id].append(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        ws_connections[session_id].remove(websocket)

async def broadcast_turn(session_id: str, turn: NegotiationTurn):
    if session_id in ws_connections:
        for ws in ws_connections[session_id]:
            try:
                await ws.send_text(turn.model_dump_json())
            except:
                pass

async def run_negotiation(session_id: str):
    session = negotiation_sessions[session_id]
    dual_env = DualAgentEnv(session)
    
    seller_runner = AgentRunner(NegotiationRole.seller, session.seller_config, openai_client, os.getenv("MODEL_NAME", "gpt-4o-mini"))
    client_runner = AgentRunner(NegotiationRole.client, session.client_config, openai_client, os.getenv("MODEL_NAME", "gpt-4o-mini"))
    
    while not dual_env.is_complete():
        # Check if human intervened
        if session.status == "paused":
            return  # exit loop cleanly; resume will spawn a new task

        # Seller turn
        obs_s = dual_env.get_observation(NegotiationRole.seller)
        action_s = await seller_runner.decide_action(obs_s)
        _, _, done = dual_env.step_seller(action_s)
        await broadcast_turn(session_id, session.negotiation_history[-1])
        await asyncio.sleep(2)
        if done: break

        # Check pause again between turns
        if session.status == "paused":
            return

        # Client turn
        obs_c = dual_env.get_observation(NegotiationRole.client)
        action_c = await client_runner.decide_action(obs_c)
        _, _, done = dual_env.step_client(action_c)
        await broadcast_turn(session_id, session.negotiation_history[-1])
        await asyncio.sleep(2)
        if done: break
        
    session.status = "completed"
    await broadcast_turn(session_id, NegotiationTurn(
        turn_number=session.turn, speaker="system", action_type="skip", clause_id="", 
        content="Negotiation complete. Please review and sign the final agreement."
    ))

STATIC_DIR = os.path.join(os.path.dirname(__file__),"..","frontend","out")
if os.path.exists(STATIC_DIR):
    app.mount("/app", StaticFiles(directory=STATIC_DIR, html=True))

@app.get("/")
def root(): 
    if os.path.exists(os.path.join(STATIC_DIR, "index.html")):
        return FileResponse(os.path.join(STATIC_DIR, "index.html"))
    return {"name":"ContractEnv","version":"1.0.0"}
